#!/usr/bin/env python3
# SPDX-License-Identifier: Apache-2.0
"""Build the launch guide: a .docx with every channel, every post, and when.

    python tools/mkkit.py --out "path/to/RESENTMENT-OS-launch-guide.docx"

Needs python-docx. The text lives here, not in a template, so the numbers
and links in it are the ones in the tree.
"""
import argparse
import datetime
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("python-docx is required: python -m pip install python-docx")

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SITE = "https://ni-sh-a-char.github.io/RESENTMENT/"
OS_URL = SITE + "os/"
REPO = "https://github.com/ni-sh-a-char/RESENTMENT"
KERNEL = "https://github.com/ni-sh-a-char/RESENTMENT---kernel"
KERNEL_SITE = "https://ni-sh-a-char.github.io/RESENTMENT---kernel/"

BRASS, CYAN, INK, MUTED = RGBColor(0x9a, 0x6a, 0x10), RGBColor(0x0d, 0x7d, 0x79), RGBColor(0x12, 0x15, 0x1c), RGBColor(0x54, 0x5c, 0x6d)

# ------------------------------------------------------------------ copy

TAGLINE = "An AI operating system on its own kernel and your own key."
HOOK = "Give an AI a computer. Take it back at a time you choose."

LINKEDIN_ORG = f"""Introducing RESENTMENT OS 2.0.0 "prahar" — an AI operating system built on our own kernel and your own key.

Most "AI operating systems" are a chat window with file access. We built one on a from-scratch kernel (x86_64, ARM64, RISC-V) and made the desktop repeat the kernel's rules at its own level:

→ Authority expires. Every permission an agent holds is a lease with a deadline. There is no "allow forever".
→ The system is a hash. One SHA-256 root over every file and setting. Snapshot it, diff it, attest it.
→ Agents are processes. Budgets, states, a process table, a kill button.
→ Nothing is forgotten. Every action is a ledger entry with a cause. Undo is a new entry, not a deletion.

Bring any key: Anthropic, OpenAI, Gemini, Groq, Mistral, xAI, DeepSeek, OpenRouter, Together, Ollama, LM Studio, or any URL you run. The key never leaves your browser. There is no server and no telemetry.

The kernel is real: a live console into ring 0 from the desktop, and CI that boots it on every push.

Open it in your browser, nothing to install: {OS_URL}
Source, Apache 2.0: {REPO}

We are honest about what is not done yet: the desktop runs in a browser and drives the kernel; running it inside the kernel needs a network stack and a compositor first. The roadmap says so out loud.

#opensource #operatingsystem #ai #agents #llm #kernel #riscv #systemsprogramming
"""

LINKEDIN_PERSONAL = f"""I built an operating system for AI agents, and today it ships under ni_sh_a.char.

The thing that bothered me about every "AI OS" I tried: the permission dialog says Allow and it means forever, and nobody can tell you what the machine looked like an hour ago. So RESENTMENT OS gives an agent a lease that expires, hashes the whole system into one number you can diff, and puts every agent in a process table with a kill button.

It runs on a kernel I wrote from scratch for three architectures. The desktop reaches it over a serial bridge because the kernel does not have a network stack yet — and I would rather say that than pretend.

Try it in your browser with your own key: {OS_URL}
Everything is Apache 2.0: {REPO}

If you work on agents, permissions, or operating systems, I would like to hear where this is wrong.
"""

X_THREAD = [
    f"{HOOK}\n\nRESENTMENT OS 2.0.0: an AI operating system on its own kernel and your own key. Runs in your browser, no install.\n\n{OS_URL}",
    "1/ Every permission an agent holds is a lease with a deadline. Five minutes, an hour, a day. There is no \"allow forever\" button. A forgotten grant stops working by itself.",
    "2/ The whole OS is one SHA-256 root. Every file, every setting, no timestamps, no keys. Snapshot, change something, snapshot, diff. Export an attestation.",
    "3/ Agents are processes. State, token budget, a row in the process table, a kill button. Killing one revokes its leases.",
    "4/ Every tool call goes in a ledger. Every file change names the call that caused it. Undo restores the files and is itself an entry. Nothing is deleted.",
    "5/ Any provider: Anthropic, OpenAI, Gemini, Groq, Mistral, xAI, DeepSeek, OpenRouter, Together, Ollama, LM Studio, or any URL. Your key stays in your browser. No server, no telemetry.",
    f"6/ Under it: a from-scratch kernel for x86_64, ARM64 and RISC-V with expiring capabilities, a Merkle root for the machine, and a transformer running on its own operators. CI boots it on every push.\n\n{KERNEL}",
    f"7/ Honest status: the desktop runs in a browser and drives the kernel over a serial bridge. Running it inside the kernel needs a network stack first. It is on the roadmap, in writing.\n\nApache 2.0. Star it: {REPO}",
]

REDDIT = [
    ("r/opensource", "I built an AI operating system where agent permissions expire and the whole system is one hash (Apache 2.0, runs in your browser)", "Text post: the four ideas in four sentences, the honest status line, both links. Reply to every comment in the first two hours."),
    ("r/osdev", "RESENTMENT: a from-scratch kernel (x86_64/ARM64/RISC-V) with expiring capabilities and a Merkle root, now with a desktop that drives it over serial", "Lead with the kernel and the bridge. This audience wants the ELF loader, SMP and the serial protocol, not the chat window. Link the kernel site first."),
    ("r/LocalLLaMA", "An AI desktop that works with Ollama/LM Studio/vLLM directly from the browser, no server, keys stay local", "Lead with local providers and OLLAMA_ORIGINS. Mention that every tool call needs a lease and shows up in a ledger."),
    ("r/selfhosted", "RESENTMENT OS: no server to host, everything in the browser, bring your own key, Apache 2.0", "Short. This audience cares that there is nothing to run and nothing phoning home."),
    ("r/programming", "An OS where the permission dialog cannot say 'forever' (design write-up)", "Link the Architecture doc, not the landing page."),
]

HN = ("Show HN: RESENTMENT OS – an AI operating system on its own kernel where permissions expire",
      f"""I built this because every AI desktop I tried had a permission dialog that meant "forever" and no way to say what the machine looked like an hour ago.

RESENTMENT OS runs in the browser on your own model key (Anthropic, OpenAI, Gemini, or any OpenAI-compatible URL including Ollama). Agents get leases instead of permissions and every lease has a deadline. The whole OS hashes to one SHA-256 root you can snapshot and diff. Every tool call is in a ledger with its cause and undo is an entry, not a deletion. Agents are processes with budgets and a kill button.

Underneath is a from-scratch kernel for x86_64, ARM64 and RISC-V ({KERNEL}) with the same ideas one level down: capabilities sealed to a time window, a Merkle DAG of every kernel object, and a transformer forward pass running on the kernel's own operators. The desktop drives it over a serial bridge; CI boots the real kernel on every push.

What is not done: the desktop runs in a browser and talks to the kernel; it does not run inside it yet, because the kernel has no network stack. That is next on its roadmap.

Static files, no server, no telemetry, Apache 2.0. Try: {OS_URL}  Source: {REPO}""")

SCHEDULE = [
    ("Day 0 (today)", "Post the LinkedIn organization post from ni_sh_a.char with the social preview image. Wait 30 minutes. Repost from your personal LinkedIn with the personal text above as commentary (not a bare repost: the commentary is what gets shown). Post the Instagram carousel with the caption. Put the site link in the Instagram bio."),
    ("Day 0, +2 h", "X thread. Pin it. Reply to yourself with a 20-second screen recording of Ctrl+K → ask → lease dialog → grant → ledger."),
    ("Day 1, morning (US time, ~14:00 IST)", "Show HN. Do not post on a weekend. Stay in the thread for three hours and answer everything, especially the criticism."),
    ("Day 1, afternoon", "r/osdev and r/LocalLLaMA (different angles, see the Reddit table). Do not post to all subreddits at once; two a day."),
    ("Day 2", "r/opensource and r/selfhosted. Instagram story: slide 3 (leases) with a 'link' sticker to the OS."),
    ("Day 3", "Dev.to / Hashnode article: 'What an OS looks like when Allow cannot mean forever' (outline below). Cross-post the link to LinkedIn as a second organization post."),
    ("Day 4", "r/programming with the architecture write-up. LinkedIn personal: a short post with one screenshot and one number that surprised you from the launch."),
    ("Day 7", "Instagram carousel two: the kernel side (reuse slides 8 and 9, add the boot transcript). LinkedIn: 'one week in' with stars, forks, the best issue, and what you changed because of feedback."),
    ("Every day, week 1", "Reply to every comment, issue and DM within a few hours. Every reply should link a doc page, not the home page. Merge the first outside pull request quickly, even a typo fix, and thank them by name."),
]

FAQ = [
    ("Is this a real OS or a web page?", "Both, honestly. The desktop is a web page today. The kernel is a real from-scratch kernel that boots on three architectures, and the desktop drives it over a serial bridge. The direction is for the desktop to run on the kernel; that needs a network stack and a compositor, which are next."),
    ("Why should I paste my API key into a website?", "You are pasting it into your own browser, not into a server. There is no RESENTMENT server. The page is static files you can read, host yourself, or run from a folder. The key goes to exactly one place: the provider you chose. You can seal it under a passphrase (AES-GCM) too."),
    ("Isn't 'authority expires' just a timeout?", "A timeout is something you add. Here it is the only kind of grant that exists: the UI cannot express 'forever'. That inverts the default, which is the whole point. The kernel does the same with cryptographic seals rather than a clock check."),
    ("Why not use the OpenAI/Anthropic SDK?", "There is no build step and no dependency by design. Three wire formats cover every provider, and the request builders and stream parsers are tested against recorded events."),
    ("Does it work offline?", "The desktop does, as an installable app. Model calls need the provider."),
    ("Why the name?", "The kernel started as a hobby project called RESENTMENT years ago, kept the name, and the OS grew on it. The codenames come from Indian time: the kernel is 'kaalachakra', the wheel of time; the OS is 'prahar', one division of the day."),
]

ARTICLE = [
    "The permission dialog that means forever (the problem, one anecdote)",
    "Leases: what changes when the UI cannot say 'always' (screenshots of the dialog and the ring)",
    "One number for the whole machine: what is hashed, what is deliberately not (timestamps, keys), a real diff",
    "Agents in the process table: budgets, waiting states, the kill button, why killing revokes leases",
    "The ledger and undo as an append (a real JSONL excerpt)",
    "The kernel underneath: the same four ideas one level down, and the serial bridge",
    "What is not done, and why the roadmap says it out loud",
    "Try it: the URL, the clone command, the three-line provider setup",
]


# ----------------------------------------------------------------- docx

def h(doc, text, lvl=1, color=None):
    p = doc.add_heading(text, level=lvl)
    if color:
        for r in p.runs:
            r.font.color.rgb = color
    return p


def para(doc, text, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def mono(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Inches(0.25)
    return p


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc, head, rows):
    t = doc.add_table(rows=1, cols=len(head))
    t.style = "Light Grid Accent 1"
    for i, c in enumerate(head):
        t.rows[0].cells[i].text = c
    for r in rows:
        cells = t.add_row().cells
        for i, c in enumerate(r):
            cells[i].text = c
    doc.add_paragraph()


def build(out, carousel_dir):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run("RESENTMENT OS 2.0.0 — launch guide")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = INK
    para(doc, TAGLINE, size=13, color=MUTED)
    para(doc, "Prepared %s for ni_sh_a.char. Everything in this document is generated from the repository by tools/mkkit.py, so the links and the claims are the ones in the tree." % datetime.date.today().isoformat(), italic=True, size=9, color=MUTED)

    h(doc, "1. The links", 1, BRASS)
    table(doc, ["What", "URL"], [
        ("Website (landing + docs)", SITE),
        ("The OS, runs in the browser", OS_URL),
        ("Source, Apache 2.0", REPO),
        ("Release v2.0.0", REPO + "/releases/tag/v2.0.0"),
        ("The kernel", KERNEL),
        ("The kernel's site", KERNEL_SITE),
        ("Maintainer", "https://github.com/PIYUSH-MISHRA-00"),
        ("Support", "https://buymeacoffee.com/piyushmishra00"),
    ])

    h(doc, "2. What is in this folder", 1, BRASS)
    table(doc, ["File", "Use it for"], [
        ("social-preview.png", "1280×640. LinkedIn link cards, X, the GitHub repository's social preview (Settings → Social preview), Dev.to cover."),
        ("carousel/slide-01.png … slide-10.png", "Instagram carousel, 1080×1350 portrait. Also works as a LinkedIn document post (upload all ten as a PDF) and as X images four at a time."),
        ("carousel/caption.txt", "The Instagram caption with hashtags. Paste as-is."),
        ("carousel/alt-text.txt", "One line per slide. Paste into each image's alt text on Instagram; it matters for reach and for accessibility."),
        ("RESENTMENT-OS-launch-guide.docx", "This document."),
    ])

    h(doc, "3. The message", 1, BRASS)
    para(doc, "One hook, one tagline, four ideas. Every post uses these words; consistency is what makes a launch look like a product rather than a weekend project.")
    para(doc, "Hook: " + HOOK, size=12)
    para(doc, "Tagline: " + TAGLINE, size=12)
    bullets(doc, [
        "Authority expires. Every permission an agent holds is a lease with a deadline. No \"allow forever\".",
        "The system is a hash. One SHA-256 root over every file and setting; snapshot, diff, attest.",
        "Agents are processes. Budgets, states, a process table, a kill button.",
        "Nothing is forgotten. Every action is a ledger entry with a cause; undo is an entry, not a deletion.",
    ])
    para(doc, "And the two lines that earn trust with a technical audience: \"Any provider, your key stays in your browser, there is no server\" and \"the desktop drives the kernel over a serial bridge; running inside it is not done yet\". Say the second one before anyone else does.")

    h(doc, "4. LinkedIn (the main channel)", 1, BRASS)
    h(doc, "4.1 Organization post, from ni_sh_a.char", 2)
    para(doc, "Attach social-preview.png as the image, or upload the ten carousel slides as a PDF document post (document posts get more dwell time on LinkedIn). Post between 9:00 and 11:00 in your audience's time zone; Tuesday to Thursday.")
    mono(doc, LINKEDIN_ORG)
    h(doc, "4.2 Personal repost, from Piyush Mishra", 2)
    para(doc, "Do not use the bare Repost button; it shows to almost nobody. Use \"Repost with your thoughts\" and paste this. Wait at least 30 minutes after the organization post so LinkedIn treats them as two posts.")
    mono(doc, LINKEDIN_PERSONAL)
    h(doc, "4.3 Follow-ups", 2)
    bullets(doc, [
        "Day 3: a second organization post linking the Dev.to article.",
        "Day 7: 'one week in' from the personal account: stars, the best issue, what changed because of feedback.",
        "Comment on every reply within a few hours; each reply is a chance for the post to be shown again.",
        "Tag nobody who has not asked to be tagged. Mention the kernel's three architectures; hardware people share.",
    ])

    h(doc, "5. Instagram", 1, BRASS)
    para(doc, "Post the ten slides as one carousel from the carousel/ folder, in order. Paste caption.txt as the caption and put each line of alt-text.txt into the matching slide's alt text (Advanced settings → Write alt text). Set the bio link to " + OS_URL + ".")
    para(doc, "Then a story: slide 03 with a link sticker to the OS and the text \"a permission that expires. swipe up.\" Reshare the post to the story on day 2 and day 5.")
    para(doc, "Reels idea (30 s, no voice): screen recording of Ctrl+K → type a question → the lease dialog appears → grant 15 min → the ledger fills → the digest in the top bar changes. Caption: \"the OS asks. you decide until when.\"")
    h(doc, "Caption", 2)
    cap = open(os.path.join(carousel_dir, "caption.txt"), encoding="utf-8").read() if os.path.exists(os.path.join(carousel_dir, "caption.txt")) else "(see carousel/caption.txt)"
    mono(doc, cap)

    h(doc, "6. X / Twitter", 1, BRASS)
    para(doc, "A thread of eight. Attach social-preview.png to the first, slides 3 to 6 to tweets 1 to 4, the boot transcript slide to tweet 6. Pin the thread.")
    for i, tw in enumerate(X_THREAD):
        mono(doc, tw)

    h(doc, "7. Hacker News", 1, BRASS)
    para(doc, "Post on a weekday, 14:00 to 16:00 IST (morning US East). Title exactly as below; HN edits titles that oversell. Post the text as the first comment, not in the URL field: submit the URL of the OS, then comment immediately.")
    para(doc, "Title: " + HN[0], size=12)
    mono(doc, HN[1])

    h(doc, "8. Reddit", 1, BRASS)
    para(doc, "Different subreddits want different angles. Two a day at most, and read each subreddit's rules on self-promotion first; r/programming wants a write-up, not a landing page.")
    table(doc, ["Subreddit", "Title", "Angle"], REDDIT)

    h(doc, "9. Article (Dev.to, Hashnode, your blog)", 1, BRASS)
    para(doc, "Title: \"What an operating system looks like when Allow cannot mean forever\". 1200 to 1800 words, screenshots from the live OS, the boot transcript from the site. Outline:")
    for i, s in enumerate(ARTICLE, 1):
        doc.add_paragraph("%d. %s" % (i, s))

    h(doc, "10. Other places worth one post each", 1, BRASS)
    table(doc, ["Where", "What"], [
        ("Product Hunt", "Launch on a Tuesday two weeks after HN, once the first issues are fixed. Tagline: the hook. Gallery: the ten slides. First comment: the HN text, shorter."),
        ("GitHub", "Set the repository's social preview to social-preview.png (Settings → Social preview). Pin the repository on the organization profile. Add the topics if missing (they are set)."),
        ("Discord / Slack communities (osdev, LocalLLaMA, Anthropic developer, r/LocalLLaMA Discord)", "One message in the showcase channel with the OS link and one sentence per idea. Never DM."),
        ("YouTube Shorts / Instagram Reels", "The 30-second recording described under Instagram. Same file, both places."),
        ("Newsletter pitches (Console.dev, TLDR, Changelog News, Hacker Newsletter)", "One email each: the hook, the OS link, the kernel link, the honest status line. They pick up projects that say what is not done."),
        ("Awesome lists", "Pull requests to awesome-selfhosted (once there is a self-host guide), awesome-osdev (the kernel), awesome-llm-apps."),
    ])

    h(doc, "11. Schedule, week one", 1, BRASS)
    table(doc, ["When", "What"], SCHEDULE)

    h(doc, "12. Questions you will get, and the answers", 1, BRASS)
    for q, a in FAQ:
        para(doc, q, size=12).runs[0].bold = True
        para(doc, a)

    h(doc, "13. Rules", 1, BRASS)
    bullets(doc, [
        "Say what is not done before anyone else does. The status table on the site exists for this; link it.",
        "Never claim the desktop runs on the kernel. It drives the kernel. The distinction is the whole credibility of the project.",
        "Every reply links a doc page, not the home page: Architecture for design questions, Providers for key questions, Kernel for the bridge, Security for trust questions.",
        "No follow-for-follow, no engagement pods, no bought reach. This audience can tell.",
        "Credit the kernel's borrowed ideas where they came from (Kaalka, WebWeaveX, SHE) when asked; the kernel README does.",
        "Merge the first outside pull request within a day, even if it is a typo.",
    ])

    h(doc, "14. What to watch", 1, BRASS)
    bullets(doc, [
        "GitHub stars and forks per day (Insights → Traffic shows referrers: you learn which channel worked).",
        "Site visits to /os/ versus the landing page: the ratio says whether people try it.",
        "Issues opened by strangers. Each one is a person who cared enough to type.",
        "Which of the four ideas people quote back. Lead with that one next time.",
    ])

    doc.save(out)
    print("  wrote %s" % out)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default=os.path.join(ROOT, "media", "RESENTMENT-OS-launch-guide.docx"))
    ap.add_argument("--carousel", default=os.path.join(ROOT, "media", "carousel"))
    args = ap.parse_args()
    os.makedirs(os.path.dirname(args.out), exist_ok=True)
    build(args.out, args.carousel)
    return 0


if __name__ == "__main__":
    sys.exit(main())
